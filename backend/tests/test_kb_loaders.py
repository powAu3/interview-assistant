from __future__ import annotations

import sys
from pathlib import Path

import pytest

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from services.kb.loaders import dispatch_loader  # noqa: E402
from services.kb.types import RawDoc  # noqa: E402


def _write(tmp: Path, name: str, content: str) -> Path:
    p = tmp / name
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(content, encoding="utf-8")
    return p


def test_dispatch_unknown_extension_returns_none(tmp_path: Path):
    p = _write(tmp_path, "a.unknown", "whatever")
    assert dispatch_loader(p) is None


def test_markdown_splits_by_headings(tmp_path: Path):
    md = """# Redis 持久化

Intro

## RDB

RDB text.

## AOF

AOF text.
"""
    p = _write(tmp_path, "redis.md", md)
    loader = dispatch_loader(p)
    assert loader is not None
    doc: RawDoc = loader.load(p, rel_path="redis.md")
    assert doc.loader == "markdown"
    assert doc.title == "Redis 持久化"
    section_paths = [s.section_path for s in doc.sections]
    assert any("RDB" in sp for sp in section_paths)
    assert any("AOF" in sp for sp in section_paths)


def test_markdown_hash_inside_code_fence_not_treated_as_heading(tmp_path: Path):
    md = "# Title\n\n```\n# not a heading\ncode line\n```\n\nAfter."
    p = _write(tmp_path, "c.md", md)
    loader = dispatch_loader(p)
    assert loader is not None
    doc = loader.load(p, rel_path="c.md")
    # "# not a heading" 不应成为新 section
    assert doc.title == "Title"
    full_text = "\n".join(s.text for s in doc.sections)
    assert "# not a heading" in full_text


def test_txt_splits_by_blank_lines(tmp_path: Path):
    content = "para1 line1\npara1 line2\n\npara2\n\npara3"
    p = _write(tmp_path, "a.txt", content)
    loader = dispatch_loader(p)
    assert loader is not None
    doc = loader.load(p, rel_path="a.txt")
    assert doc.loader == "txt"
    assert len(doc.sections) == 3


def test_empty_file_returns_doc_with_no_sections(tmp_path: Path):
    p = _write(tmp_path, "empty.md", "")
    loader = dispatch_loader(p)
    doc = loader.load(p, rel_path="empty.md")
    assert doc.sections == []


def test_docx_headings_become_sections(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_heading("Redlock", level=1)
    doc.add_paragraph("intro text")
    doc.add_heading("Pros", level=2)
    doc.add_paragraph("p1")
    doc.add_heading("Cons", level=2)
    doc.add_paragraph("c1")
    p = tmp_path / "redlock.docx"
    doc.save(str(p))

    loader = dispatch_loader(p)
    assert loader is not None
    raw = loader.load(p, rel_path="redlock.docx")
    assert raw.loader == "docx"
    assert raw.title == "Redlock"
    paths = [s.section_path for s in raw.sections]
    assert any("Pros" in sp for sp in paths)
    assert any("Cons" in sp for sp in paths)


def test_pdf_l1_extracts_text_per_page(tmp_path: Path):
    pytest.importorskip("pypdf")
    pytest.importorskip("reportlab.pdfgen.canvas")
    from reportlab.pdfgen import canvas

    p = tmp_path / "notes.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(100, 750, "Page one: Redis persistence")
    c.showPage()
    c.drawString(100, 750, "Page two: Go GMP")
    c.showPage()
    c.save()

    loader = dispatch_loader(p)
    assert loader is not None
    raw = loader.load(p, rel_path="notes.pdf")
    assert raw.loader == "pdf"
    texts = [s.text for s in raw.sections]
    assert any("Redis" in t for t in texts)
    assert any("GMP" in t for t in texts)
    pages = [s.page for s in raw.sections]
    assert pages == [1, 2]


def test_pdf_l2_ocr_stubbed(monkeypatch, tmp_path: Path):
    """用 stub 模拟 OCR，避免依赖 rapidocr 模型。"""
    pytest.importorskip("pypdf")
    pytest.importorskip("reportlab.pdfgen.canvas")
    from reportlab.pdfgen import canvas

    p = tmp_path / "scan.pdf"
    c = canvas.Canvas(str(p))
    # 只画 3 个字符，文本层极短，触发 OCR 分支
    c.drawString(100, 750, "abc")
    c.save()

    from services.kb.loaders import pdf as pdf_loader
    monkeypatch.setattr(
        pdf_loader, "_run_ocr_on_pdf_page", lambda pdf_path, idx: "OCR TEXT FROM IMAGE"
    )
    monkeypatch.setattr(pdf_loader, "_l2_ocr_enabled", lambda: True)
    monkeypatch.setattr(pdf_loader, "_l3_vision_enabled", lambda: False)

    loader = dispatch_loader(p)
    raw = loader.load(p, rel_path="scan.pdf")
    ocr_sections = [s for s in raw.sections if s.origin == "ocr"]
    assert ocr_sections
    assert "OCR TEXT" in ocr_sections[0].text
    assert ocr_sections[0].section_path.endswith("(OCR)")


def test_pdf_l3_vision_stubbed(monkeypatch, tmp_path: Path):
    """用 stub 模拟 Vision caption。"""
    pytest.importorskip("pypdf")
    pytest.importorskip("reportlab.pdfgen.canvas")
    from reportlab.pdfgen import canvas

    p = tmp_path / "diagram.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(100, 750, "some text")
    c.save()

    from services.kb.loaders import pdf as pdf_loader
    monkeypatch.setattr(
        pdf_loader, "_run_vision_caption", lambda pdf_path, idx: "架构图: A → B → C"
    )
    monkeypatch.setattr(pdf_loader, "_l2_ocr_enabled", lambda: False)
    monkeypatch.setattr(pdf_loader, "_l3_vision_enabled", lambda: True)

    loader = dispatch_loader(p)
    raw = loader.load(p, rel_path="diagram.pdf")
    vision_sections = [s for s in raw.sections if s.origin == "vision"]
    assert vision_sections
    assert "架构图" in vision_sections[0].text
