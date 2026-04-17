"""DOCX loader：按 Heading 1..9 层级切 section，表格以 tab 分隔附在末尾 section。"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from ..types import RawDoc, RawSection
from ._base import register

_log = logging.getLogger(__name__)

try:
    from docx import Document  # python-docx
    _DOCX_OK = True
except Exception as e:  # pragma: no cover
    _log.info("docx loader not available: %s", e)
    _DOCX_OK = False


class DocxLoader:
    name = "docx"
    extensions = (".docx",)

    def load(self, file_path: Path, *, rel_path: str) -> RawDoc:
        if not _DOCX_OK:
            raise RuntimeError("python-docx 未安装：请 `pip install python-docx`")
        doc = Document(str(file_path))
        title: Optional[str] = None
        heading_stack: list[tuple[int, str]] = []
        sections: list[RawSection] = []
        buf: list[str] = []
        current_sp = ""

        def flush() -> None:
            nonlocal buf
            if buf:
                txt = "\n".join(buf).strip()
                if txt:
                    sections.append(RawSection(section_path=current_sp, text=txt))
                buf = []

        for para in doc.paragraphs:
            style = (para.style.name or "").strip() if para.style else ""
            text = (para.text or "").rstrip()
            level = _heading_level(style)
            if level is not None and text:
                if title is None and level == 1:
                    title = text
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text))
                flush()
                current_sp = _format_sp(title, heading_stack)
                continue
            if text:
                buf.append(text)

        for tbl in doc.tables:
            rows: list[str] = []
            for row in tbl.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append("\t".join(cells))
            if rows:
                buf.append("\n".join(rows))
        flush()
        return RawDoc(
            path=rel_path,
            title=title or Path(rel_path).stem,
            loader="docx",
            sections=sections,
        )


def _heading_level(style: str) -> Optional[int]:
    if not style:
        return None
    low = style.lower()
    if low.startswith("heading"):
        digits = "".join(ch for ch in style if ch.isdigit())
        if digits:
            try:
                lv = int(digits)
                if 1 <= lv <= 9:
                    return lv
            except ValueError:
                return None
    return None


def _format_sp(title: Optional[str], stack: list[tuple[int, str]]) -> str:
    parts: list[str] = []
    if title:
        parts.append(title)
    for _, h in stack:
        if h and (not parts or parts[-1] != h):
            parts.append(h)
    return " > ".join(parts)


if _DOCX_OK:
    register(DocxLoader())

__all__ = ["DocxLoader"]
